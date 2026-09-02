import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_styled_document():
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    return doc

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy Blue
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 118, 110) # Teal Blue
    return h

def add_p(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(15, 23, 42)
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    return p

# Ensure docs/ directory exists as canonical location
DOCS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../docs"))
os.makedirs(DOCS_DIR, exist_ok=True)

# ==============================================================================
# 1. GENERATE TECHNICAL DOCUMENTATION (.docx & .txt)
# ==============================================================================

def generate_technical_docs():
    print("Generating TECHNICAL_DOCUMENTATION.docx & .txt...")
    doc = create_styled_document()
    
    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run("MALE UAV Aero Piston Engine Digital Twin")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(30, 58, 138)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Comprehensive System Architecture & Technical Specifications\nDRDO / SIH Defense Hackathon Project")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    add_heading_1(doc, "1. Executive Summary")
    add_p(doc, "The Medium-Altitude Long-Endurance (MALE) UAV Aero Piston Engine Digital Twin is a defense-grade, 5-service FastAPI microservices platform designed for real-time telemetry monitoring, predictive maintenance, physics-informed anomaly detection, and explainable AI (XAI) diagnostics.")
    add_p(doc, "The system processes high-frequency Engine Control Unit (ECU) telemetries via a simulated CAN bus (engine_can.dbc), engineers 131 temporal and statistical feature signals in real-time, executes 4 machine learning models with cryptographic SHA-256 integrity verification, and persists mission replays and maintenance advisories to MongoDB Atlas.")

    add_heading_1(doc, "2. System Architecture & Microservices Topology")
    add_p(doc, "The platform adopts a decoupled, 5-tier microservices architecture designed for scalability, fault isolation, and horizontal expansion across defense ground control stations.")

    # Microservices Table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["Service Name", "Port", "Core Responsibilities", "Security Control"]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    services_info = [
        ("API Gateway", "8000", "Central ingress, WebSocket stream (/ws/telemetry), CORS, IP rate limiting (60 req/min), internal request proxying.", "Sliding Window Limiter & HTTPS"),
        ("Telemetry Service", "8001", "Mission simulation engine, CAN bus encoding/decoding (DBC parser), baseline physics models, 131 feature engineering, Scenario simulator.", "X-Internal-Key Enforcement"),
        ("ML Inference Service", "8002", "Execution of 4 AI models (Anomaly, Degradation, Fault, RUL) with model artifact SHA-256 hash verification.", "X-Internal-Key & Hash Verification"),
        ("XAI & Advisory Service", "8003", "SHAP TreeExplainer feature attributions, top diagnostic drivers, state-tracked anti-spam maintenance advisories.", "X-Internal-Key Enforcement"),
        ("MongoDB Atlas Service", "8004", "Async persistence proxy for MongoDB Atlas cloud (frame telemetry, mission replays, advisory audit history).", "X-Internal-Key & TLS Connection")
    ]
    
    for row_idx, row_data in enumerate(services_info):
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            set_cell_background(row_cells[i], "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(30, 41, 59)

    add_heading_1(doc, "3. Defense AI/ML Model Suite")
    add_p(doc, "The predictive analytics engine combines 4 machine learning models trained on aero piston engine operational endurance datasets:")
    add_bullet(doc, "Uses decision score thresholding on 13 baseline sensor signals to detect operating envelope deviations.", "1. Anomaly Detection (Isolation Forest): ")
    add_bullet(doc, "Predicts engine Health Index % (0-100%) and continuous degradation index from 120 rolling features.", "2. Degradation Estimation (XGBoost Regressor): ")
    add_bullet(doc, "Classifies failure modes across 6 classes: injector_degradation, lubrication_degradation, misfire, normal, overheating, sensor_fault.", "3. Fault Classification (Multiclass XGBoost): ")
    add_bullet(doc, "Predicts Remaining Useful Life in flight hours using an XGBoost Regressor (131 features, with post-processing: degradation-aware anchoring, exponential smoothing, slew-rate limiting, and quantile-sampling-based uncertainty quantification for P10-P90 confidence intervals).", "4. RUL Prediction (XGBoost Regressor): ")

    add_heading_1(doc, "4. Defense-in-Depth Security Framework")
    add_p(doc, "The platform incorporates five security control layers enforcing zero-trust principles across inter-service communication:")
    add_bullet(doc, "Only Port 8000 is exposed to external GCS clients. Microservices on ports 8001-8004 validate internal access keys via HTTP headers.", "1. Service-to-Service Authentication (X-Internal-Key): ")
    add_bullet(doc, "API Gateway limits client requests per IP address (60 requests/minute) returning HTTP 429 on abuse.", "2. API Gateway Rate Limiting: ")
    add_bullet(doc, "CAN codec appends and verifies a 1-byte XOR checksum on binary CAN frames, logging audit warnings on payload tampering.", "3. CAN Frame Checksum Integrity: ")
    add_bullet(doc, "Model loader computes SHA-256 hashes of all model artifacts at startup, comparing them against models/model_hashes.json.", "4. AI Model Hash Integrity & Traceability: ")
    add_bullet(doc, "Production deployment supports RSA 2048-bit SSL certificates for HTTPS/WSS encryption.", "5. Transport Layer Security (TLS 1.3): ")

    # Save .docx into docs/ ONLY
    docx_path = os.path.join(DOCS_DIR, "TECHNICAL_DOCUMENTATION.docx")
    doc.save(docx_path)
    print(f"Saved {docx_path}")

    # Generate .txt version into docs/ ONLY
    txt_path = os.path.join(DOCS_DIR, "TECHNICAL_DOCUMENTATION.txt")
    txt_content = """================================================================================
MALE UAV AERO PISTON ENGINE DIGITAL TWIN — TECHNICAL DOCUMENTATION
DRDO / SIH Defense Hackathon Project Architecture & Security Specifications
================================================================================

1. EXECUTIVE SUMMARY
The Medium-Altitude Long-Endurance (MALE) UAV Aero Piston Engine Digital Twin is a defense-grade, 5-service FastAPI microservices platform designed for real-time telemetry monitoring, predictive maintenance, physics-informed anomaly detection, and explainable AI (XAI) diagnostics.

2. SYSTEM ARCHITECTURE & MICROSERVICES TOPOLOGY
  - API Gateway (Port 8000): Ingress proxy, WebSocket stream (/ws/telemetry), Rate Limiting.
  - Telemetry Service (Port 8001): Mission simulation engine, CAN bus parser, 131 feature engineering, Scenario Simulator.
  - ML Inference Service (Port 8002): Anomaly, Degradation, Fault, RUL models with SHA-256 verification.
  - XAI & Advisory Service (Port 8003): SHAP TreeExplainer, top diagnostic drivers, advisories.
  - MongoDB Atlas Service (Port 8004): Async MongoDB cloud persistence proxy.

3. DEFENSE AI/ML MODEL SUITE
  1. Isolation Forest: Anomaly Detection
  2. XGBoost Regressor: Engine Degradation Index & Health %
  3. Multiclass XGBoost: Fault Classifier (6 classes: injector_degradation, lubrication_degradation, misfire, normal, overheating, sensor_fault)
  4. XGBoost Regressor (131 features, with post-processing: degradation-aware anchoring, exponential smoothing, slew-rate limiting, and quantile-sampling-based uncertainty quantification for P10-P90 confidence intervals): RUL Prediction

4. DEFENSE-IN-DEPTH SECURITY FRAMEWORK
  1. Service-to-Service Auth (X-Internal-Key)
  2. API Gateway Rate Limiting (60 req/min)
  3. CAN Bus Frame Checksum Integrity (1-byte XOR)
  4. SHA-256 AI Model Fingerprinting & Hash Verification
  5. Transport Layer Security (TLS 1.3 / HTTPS / WSS)
"""
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(txt_content)
    print(f"Saved {txt_path}")

# ==============================================================================
# 2. GENERATE DEPLOYMENT ROADMAP (.docx & .txt)
# ==============================================================================

def generate_deployment_roadmap():
    print("Generating DEPLOYMENT_ROADMAP.docx & .txt...")
    doc = create_styled_document()
    
    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run("MALE UAV Aero Piston Engine Digital Twin")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(30, 58, 138)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("Production Deployment & Enterprise Defense Integration Roadmap\nPhased Transition Matrix (MVP -> Defense Standard)")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    add_heading_1(doc, "1. Executive Summary & Strategic Objectives")
    add_p(doc, "This roadmap defines the structured engineering transition from the current hackathon MVP baseline to a fully certified, air-gapped, defense-grade digital twin platform for MALE UAV operations in accordance with DRDO and Indian Armed Forces operational standards.")

    add_heading_1(doc, "2. Phased Deployment Timeline & Milestone Matrix")

    # Roadmap Table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["Phase & Timeline", "Key Milestones", "Core Deliverables", "Target Environment"]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    phases_info = [
        ("Phase 1: MVP Baseline (Current / Complete)", "5 microservices operational, MongoDB Atlas logging, SHA-256 model hashes, CAN XOR checksums, Scenario Simulator.", "Local multi-service launcher, FastAPI Swagger docs, 11-test verification script.", "Local Dev / Demo Sandbox"),
        ("Phase 2: Short-Term (0 - 3 Months)", "Dockerization, Kubernetes Helm charts, OAuth2 / JWT Role-Based Access Control (RBAC).", "Docker Compose configuration, CI/CD pipeline via GitHub Actions, Prometheus + Grafana dashboards.", "Staging / Cloud Testbed"),
        ("Phase 3: Medium-Term (3 - 9 Months)", "Hardware-in-the-Loop (HIL) SocketCAN integration, AUTOSAR SecOC implementation, mTLS.", "Peak-CAN hardware interface drivers, mTLS client certs, air-gapped local MongoDB Enterprise cluster.", "Ground Test Rig / HIL Lab"),
        ("Phase 4: Defense Production (9 - 18 Months)", "Hardware Security Module (HSM / TPM 2.0) keys, Edge AI quantization (TensorRT), MIL-STD-810H cert.", "Airborne flight computer binary, Deep Learning Bus IDS, DRDO defense certification compliance.", "Operational UAV GCS & Flight Unit")
    ]
    
    for row_idx, row_data in enumerate(phases_info):
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            set_cell_background(row_cells[i], "F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(30, 41, 59)

    # Save .docx into docs/ ONLY
    docx_path = os.path.join(DOCS_DIR, "DEPLOYMENT_ROADMAP.docx")
    doc.save(docx_path)
    print(f"Saved {docx_path}")

    # Generate .txt version into docs/ ONLY
    txt_path = os.path.join(DOCS_DIR, "DEPLOYMENT_ROADMAP.txt")
    txt_content = """================================================================================
MALE UAV AERO PISTON ENGINE DIGITAL TWIN — PRODUCTION DEPLOYMENT ROADMAP
Phased Enterprise & Defense Transition Plan (DRDO / SIH Standards)
================================================================================

1. EXECUTIVE SUMMARY
This roadmap defines the structured engineering transition from the current hackathon MVP baseline to a fully certified, air-gapped, defense-grade digital twin platform for MALE UAV operations.

2. PHASED TIMELINE & MILESTONES

[PHASE 1: HACKATHON MVP BASELINE — COMPLETED]
  - 5 microservices operational (Ports 8000-8004)
  - MongoDB Atlas cloud integration & mission replay
  - SHA-256 model hash integrity & 131-feature RUL alignment
  - Defense-in-depth security layer (X-Internal-Key, rate limiter, CAN XOR checksums)
  - Scenario simulator (/api/simulation/scenario) & vibration/coking advisories

[PHASE 2: ENTERPRISE CONTAINERIZATION & RBAC — 0 to 3 MONTHS]
  - Docker & Kubernetes Helm charts for all 5 services
  - OAuth2 / JWT Role-Based Access Control (Pilot vs Tech vs Admin)
  - Prometheus + Grafana real-time monitoring dashboard
  - GitHub Actions CI/CD automated deployment pipeline

[PHASE 3: HARDWARE-IN-THE-LOOP & AUTOSAR SecOC — 3 to 9 MONTHS]
  - Peak-CAN / Vector physical SocketCAN hardware interface
  - AUTOSAR SecOC (Secure On-Board Communication) for CAN-FD frames
  - Mutual TLS (mTLS) with X.509 client certificates between microservices
  - Air-gapped local MongoDB Enterprise cluster for classified environments

[PHASE 4: FLIGHT COMPUTER QUANTIZATION & MIL-STD CERTIFICATION — 9 to 18 MONTHS]
  - TPM 2.0 / HSM Hardware Security Module key storage
  - TensorRT / ONNX Runtime AI model quantization for airborne flight computers
  - Deep Learning CAN Bus Intrusion Detection System (BIDS)
  - DRDO / MIL-STD-810H environmental & software reliability certification
"""
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(txt_content)
    print(f"Saved {txt_path}")

if __name__ == "__main__":
    generate_technical_docs()
    generate_deployment_roadmap()
